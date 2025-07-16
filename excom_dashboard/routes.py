from flask import render_template, request, redirect, url_for, session
from . import excom_bp
import os
import csv
from datetime import datetime
from docx import Document

UPLOAD_FOLDER = 'uploads'
GENERATED_FOLDER = 'generated'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(GENERATED_FOLDER, exist_ok=True)

@excom_bp.route('/email', methods=['GET', 'POST'], endpoint='bulk_email')
def bulk_email():
    message = ''
    user = session.get('user', 'guest')
    role = session.get('role', 'member')

    if request.method == 'POST':
        subject = request.form['subject']
        body = request.form['message']
        csv_file = request.files.get('csv')

        if csv_file and csv_file.filename.endswith('.csv'):
            path = os.path.join(UPLOAD_FOLDER, csv_file.filename)
            csv_file.save(path)
            # Simulated email logic here
            message = f"Email list '{csv_file.filename}' uploaded. Emails simulated."
        else:
            message = "Please upload a valid CSV file."

    return render_template(
        'bulk_email.html',
        user=user,
        role=role,
        message=message
    )


@excom_bp.route('/mom', methods=['GET', 'POST'],endpoint='mom_generator')
def mom_generator():
    message = ''
    user = session.get('user', 'guest')
    role = session.get('role', 'member')

    if request.method == 'POST':
        agenda = request.form['agenda']
        date = request.form['date'] or datetime.now().strftime('%B %d, %Y')
        attendees = request.form['attendees']
        discussion = request.form['discussion']

        doc = Document()
        doc.add_heading('Minutes of Meeting', 0)
        doc.add_paragraph(f"Title: {agenda}")
        doc.add_paragraph(f"Date: {date}")
        doc.add_paragraph(f"Attendees: {attendees}")
        doc.add_paragraph("Discussion Points:")

        for point in discussion.splitlines():
            doc.add_paragraph(f"- {point}", style='List Bullet')

        filename = f"MOM_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(GENERATED_FOLDER, filename)
        doc.save(filepath)

        return redirect(url_for('static', filename=f'../{GENERATED_FOLDER}/{filename}'))

    return render_template(
        'mom_generator.html',
        user=user,
        role=role,
        message=message
    )