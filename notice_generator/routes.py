# === notice_generator/routes.py ===
from flask import Blueprint, render_template, request, send_file, session
from docx import Document
from docx.shared import Inches
import os
from datetime import datetime

notice_bp = Blueprint('notice', __name__, template_folder='templates')

UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "generated"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

@notice_bp.route('/notice', methods=['GET', 'POST'], endpoint='generate_notice')
def generate_notice():
    message = None
    user = session.get('user', 'guest')
    role = session.get('role', 'member')

    if request.method == 'POST':
        subject     = request.form['subject']
        recipient   = request.form['recipient']
        body        = request.form['body']
        your_name   = request.form['your_name']
        your_pos    = request.form['your_pos']
        date        = request.form.get('date') or datetime.now().strftime('%B %d, %Y')

        logo        = request.files.get('logo')
        signature   = request.files.get('signature')

        logo_path = os.path.join(UPLOAD_FOLDER, logo.filename) if logo and logo.filename else None
        sig_path  = os.path.join(UPLOAD_FOLDER, signature.filename) if signature and signature.filename else None

        if logo_path: logo.save(logo_path)
        if sig_path: signature.save(sig_path)

        notice_number = f"NOTICE-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
        filename = f"{notice_number}.docx"
        output_path = os.path.join(OUTPUT_FOLDER, filename)

        # === Document Creation ===
        doc = Document()

        if logo_path and os.path.exists(logo_path):
            doc.add_picture(logo_path, width=Inches(1.2))

        doc.add_heading("NOTICE", level=0)
        doc.add_paragraph("IEEE VSSUT STUDENT BRANCH")
        doc.add_paragraph("Burla, Sambalpur - 768018, Odisha")
        doc.add_paragraph("Email: ieee@vssut.ac.in")

        doc.add_paragraph(f"\nNotice No: {notice_number}")
        doc.add_paragraph(f"Date: {date}")
        doc.add_paragraph(f"\nCopy To: {recipient}")

        doc.add_heading(f"Subject: {subject}", level=1)
        doc.add_paragraph("Dear Sir/Madam,")
        doc.add_paragraph(body)
        doc.add_paragraph("\nThank you.")
        doc.add_paragraph("\nSincerely,")
        doc.add_paragraph(your_name)
        doc.add_paragraph(your_pos)

        if sig_path and os.path.exists(sig_path):
            doc.add_picture(sig_path, width=Inches(1.5))

        doc.save(output_path)

        return send_file(output_path, as_attachment=True)

    return render_template("notice.html", user=user, role=role, message=message)
